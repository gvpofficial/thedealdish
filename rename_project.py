import os
import docx
import win32com.client
import subprocess

def replace_in_file(filepath, replacements):
    """Replaces text in a standard text file."""
    with open(filepath, 'r', encoding='utf-8') as file:
        content = file.read()
    
    modified = content
    for old, new in replacements.items():
        modified = modified.replace(old, new)
        
    if modified != content:
        with open(filepath, 'w', encoding='utf-8') as file:
            file.write(modified)
        print(f"Updated text in {os.path.basename(filepath)}")

def replace_text_in_paragraph(p, old_text, new_text):
    """Replaces text inside docx paragraph preserving style runs."""
    if old_text in p.text:
        # First try replacing in individual runs
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        
        # If it was split across runs, merge them or do a paragraph-level replace
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)

def replace_text_in_docx(doc, old_text, new_text):
    """Scans and replaces text in paragraphs, tables, headers, and footers."""
    # 1. Main body paragraphs
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, old_text, new_text)
        
    # 2. Main body tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, old_text, new_text)
                    
    # 3. Headers and footers of all sections
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_text_in_paragraph(p, old_text, new_text)
        for p in section.footer.paragraphs:
            replace_text_in_paragraph(p, old_text, new_text)

def update_docx_project_name(docx_path):
    """Opens docx, performs case replacements, and saves."""
    doc = docx.Document(docx_path)
    
    # 1. UPPER CASE DEALDISH -> THEDEALDISH
    replace_text_docx_recursive(doc, "DEALDISH", "THEDEALDISH")
    
    # 2. CamelCase DealDish -> TheDealDish
    replace_text_docx_recursive(doc, "DealDish", "TheDealDish")
    
    doc.save(docx_path)
    print("Updated text inside Word Document.")

def replace_text_docx_recursive(doc, old_text, new_text):
    # Wrapper for recursive replacement
    replace_text_in_docx(doc, old_text, new_text)

def main():
    workspace_dir = r"D:\gvp codes\deal-dish"
    
    # 1. Text File Replacements (Website assets)
    web_replacements = {
        "DealDish": "TheDealDish",
        "DEALDISH": "THEDEALDISH",
        "dealdish": "thedealdish"
    }
    
    replace_in_file(os.path.join(workspace_dir, "index.html"), web_replacements)
    replace_in_file(os.path.join(workspace_dir, "app.js"), web_replacements)
    replace_in_file(os.path.join(workspace_dir, "style.css"), web_replacements)
    
    # 2. Word Document Replacements
    old_docx = os.path.join(workspace_dir, "DealDish_Synopsis_Redesigned.docx")
    new_docx = os.path.join(workspace_dir, "THEDEALDISH_Synopsis_Redesigned.docx")
    
    if os.path.exists(old_docx):
        update_docx_project_name(old_docx)
        
        # Rename DOCX in filesystem
        if os.path.exists(new_docx):
            os.remove(new_docx)
        os.rename(old_docx, new_docx)
        print("Renamed DOCX file to THEDEALDISH_Synopsis_Redesigned.docx")
    else:
        print("Original DOCX file not found!")
        
    # 3. PDF Compilation via Word COM
    new_pdf = os.path.join(workspace_dir, "THEDEALDISH_Synopsis.pdf")
    old_pdf = os.path.join(workspace_dir, "DealDish_Synopsis.pdf")
    
    if os.path.exists(new_docx):
        print("Initializing Word COM interface for PDF compilation...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            abs_doc = os.path.abspath(new_docx)
            abs_pdf = os.path.abspath(new_pdf)
            wdoc = word.Documents.Open(abs_doc)
            wdoc.SaveAs(abs_pdf, FileFormat=17) # 17 is wdFormatPDF
            wdoc.Close()
            print("Successfully compiled updated PDF: THEDEALDISH_Synopsis.pdf")
            
            # Clean up old files if they exist
            if os.path.exists(old_pdf):
                os.remove(old_pdf)
                print("Removed old PDF file.")
        except Exception as e:
            print(f"Failed to convert PDF via Word COM: {e}")
        finally:
            word.Quit()
            
    # 4. Stage changes in Git
    try:
        subprocess.run(["git", "add", "index.html", "app.js", "style.css"], cwd=workspace_dir, check=True)
        # Git untrack and remove old files
        if os.path.exists(new_docx):
            subprocess.run(["git", "add", "THEDEALDISH_Synopsis_Redesigned.docx", "THEDEALDISH_Synopsis.pdf"], cwd=workspace_dir)
        subprocess.run(["git", "rm", "DealDish_Synopsis_Redesigned.docx", "DealDish_Synopsis.pdf"], cwd=workspace_dir, stderr=subprocess.DEVNULL)
        print("Staged all changes in Git.")
    except Exception as e:
        print(f"Git staging warning: {e}")

if __name__ == "__main__":
    main()
